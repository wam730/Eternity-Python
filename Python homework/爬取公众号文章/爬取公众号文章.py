import requests
from time import sleep
from os import mkdir
from os.path import isdir
from bs4 import BeautifulSoup
from docx import Document, opc, oxml
from docx.shared import Inches

dstDir = 'D:\\source\\Word'
if not isdir(dstDir):
    mkdir(dstDir)

url = r'https://mp.weixin.qq.com/s/u9FeqoBaA3Mr0fPCUMbpqA'
content = requests.get(url)
content.encoding = 'utf8'
soupMain = BeautifulSoup(content.text,'lxml')

for a in soupMain.find_all('a',target='_blank'):
    #sleep(1)
    title = a.text.replace('\\','').replace('|','_').replace('/','_')
    print(title)
    link = a['href']
    currentDocument = Document()
    currentDocument.add_heading(title)
    content = requests.get(link)
    content.encoding = 'utf8'
    soup = BeautifulSoup(content.text,'lxml').find('div',
                                                   id = 'js_content')
    if not soup:
        continue
    for  child in soup.children:
        child = BeautifulSoup(str(child),'lxml')
        if child.a:
            p = currentDocument.add_paragraph(text=child.text)
            try:
                p.add_run()
                r_id = p.part.relate_to(child.a['href'],
                                        opc.constants.RELATIONSHIP_TYPE.HYPERLIN,is_external = True)
                hyperlink = oxml.shared.OxmlElement('w:hyperlink')
                hyperlink.set(oxml.shared.qn('r:id'),r_id)
                hyperlink.append(p.runs[0]._r)
            except:
                pass
        elif child.img:
            pic='temp.png'
            with open(pic,'wb') as fp:
                fp.write(requests.get(child.img['data-src']).content)
            try:
                currentDocument.add_picture(pic,width=Inches(4))
            except:
                pass
        elif child.tr:
            rows = child.find_all('tr')
            cols = rows[0].find_all('td')
            table = currentDocument.add_table(len(rows),len(cols))
            for rindex, row in enumerate(rows):
                for cindex, col in enumerate(row.find_all('td')):
                    try:
                        cell = table.cell(rindex,cindex)
                        cell.text = col.text
                    except:
                        pass
        else:
            para = child.text
            currentDocument.add_paragraph(text=para)
    currentDocument.save(dstDir+'\\'+title+'.docx')
